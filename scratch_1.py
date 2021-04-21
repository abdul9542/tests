# from docxtpl import DocxTemplate
#
# def loop(n):
#     doc = DocxTemplate("sample.txt")
#     context= {1:{"title":"Alfie", "date":"2020-03-03"},
#               2:{"title":"Derek", "date":"2020-03-03"},
#               3:{"title":"Alfie", "date":"2020-03-03"}}
#     doc.render(context[n])
#     doc.save("%s.docx" %str(n))
# # doc.render(context)
# # doc.save("generated_doc.docx")
#
# for i in range(1,4):
#     loop(i)


import docx
import os
from docx.shared import Inches
doc = docx.Document()
doc.add_heading("Test Doc", 0)
parag=doc.add_paragraph("hello")
doc.add_heading("Heading 1 ",1)
parag.add_run("u r in ms office")
parag.add_run("Abdul").bold=True
doc.add_heading("Heading 2",2)
italicparag=doc.add_paragraph()
italicparag.add_run("ïtallic fun").italic=True
doc.add_picture('123.png',width=Inches(2.8))
doc.save("test.docx")
os.system("start test.docx")